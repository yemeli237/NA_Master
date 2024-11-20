import tkinter
from typing import *
import customtkinter
from tkinter import *
from docx import Document
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
# import filtre
import os
import json
from tkinter import filedialog
import time
import webbrowser
import dispo
import detail 
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import docx
import sys
from  docx.api import Document
from docx.shared import Pt, RGBColor 
from docx.oxml.ns import qn
import asyncio

from reportlab.lib.pagesizes import letter 
from reportlab.pdfgen import canvas
import io
# from fpdf import FPDF






customtkinter.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"

class App(customtkinter.CTk, tkinter.Tk, ):
    def __init__(self):
        super().__init__()
        self.liste = []
        self.source = ""
        self.auto()
        temps = 0
        # pdf = canvas.Canvas("exemple.pdf",pagesize=letter)
        
        # try:
        #     if os.path.exists("filtre.py"):
        #         liste = filtre.filtre()
        #         print(liste["Alarm ID"])
        # except:
        #     print("un problem c'est produit")
                
            
        
        #configuration de l'ecrant principal
        
        
        self.title("NA_Master")
        self.geometry(f"{1080}x{720}+{150}+{0}")
        self.minsize(width=1080, height=720)
        self.iconbitmap("icon.ico")
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure((2, 3), weight=0)
        self.grid_rowconfigure((0, 1, 2), weight=1)
        
        
        
        
        #configuration de la navbar de gauche
        self.navbar =  customtkinter.CTkFrame(
            self,
            width=150,
            corner_radius=0,
        )
        self.navbar.grid(row=0, column=0, sticky="nsew",  
               rowspan=4
               )
        self.navbar.grid_rowconfigure(5, weight =1)
        
        #nom de l'aplication
        self.nom_app = customtkinter.CTkLabel(
            self.navbar,
            text="NA_Master",
            font=customtkinter.CTkFont(size=25, weight="bold")
        ).grid(row=0, column=0,padx=20, pady=(20, 10))
        
        #choix du site
        self.site_label = customtkinter.CTkLabel(
            self.navbar,
            text="site a superviser",
            anchor="w"
        ).grid(row=1, column=0, padx =20, pady=(10, 0))
        self.site_frame = customtkinter.CTkOptionMenu(
            self.navbar,
            values=[
                "Mfoundi",
                "Hors Mfoundi"
            ],
            command=self.choix_site
        ).grid(row=2, column=0,padx =20, pady=(0, 10))
        
        #########################################################
        #recapitulatif
        self.recap_label = customtkinter.CTkLabel(
            self.navbar,
            text="Disponibilite",
            anchor="w"
        ).grid(row=3, column=0,)
        self.recap_frame = customtkinter.CTkOptionMenu(
            self.navbar,
            values=[
                "Mfoundi",
                "Hors Mfoundi"
            ],
            command=self.choix_recap
        ).grid(row=4, column=0,padx =20, pady=(0, 10))
        
        
        #############################################
        #actualiser
        # self.actu = customtkinter.CTkLabel(master=self.navbar, text="Actualiser").grid(row=5, column=0, padx =20, pady=(0, 10))
        self.act_btn = customtkinter.CTkButton(
            master=self.navbar,
            text="Actualiser",
            command=self.auto,
        ).grid(row=5, column=0, padx =20, pady=(0, 10))
        
        ####################################################
        #choix du theme
        self.theme_label = customtkinter.CTkLabel(
            self.navbar,
            text="choix du theme",
            anchor="w"
        ).grid(row=6, column=0,padx =20, pady=(0, 10))
        self.theme = customtkinter.CTkOptionMenu(
            self.navbar,
            values=[
               "Light",
               "Dark",
               "System"
            ],
            command=self.choix_theme
        ).grid(row=7,  column=0,padx =20, pady=(0, 10))
        
        ################################################
        #echelle
        self.theme_label = customtkinter.CTkLabel(
            self.navbar,
            text="echelle",
            anchor="w"
        ).grid(row=8, column=0,)
        self.echelle = customtkinter.CTkOptionMenu(
            self.navbar,
            values=[
               "80%",
               "90%",
               "100%",
               "110%",
               "120%",
            ],
            command=self.choix_echel
        )
        self.echelle.grid(row=9, column=0,padx =20, pady=(0, 10))
        self.echelle.set("100%")
        
        
        ########################################################
        #configuration
        self.bar_config = customtkinter.CTkFrame(master=self, height=50, border_color="#1492E6", border_width=1)
        self.bar_config.grid(row=3, column=1, columnspan=2, padx=(20, 0), pady=(0, 0), sticky="nsew")
        
        self.source_label = customtkinter.CTkOptionMenu(master=self.bar_config,values=[
            "Source",
            "Stokage Interne",
            "Internet",
            ],
            command=self.sources)
        self.source_label.grid(row=0, column=0, pady=(5, 5), padx=(5, 0))
        self.lacal_source = customtkinter.CTkButton(master=self, text="Valide")
        self.lacal_source.grid(row=4, column=3, pady=(5, 5))
        # self.lacal_internet = customtkinter.CTkButton(master=self.bar_config, text="Internet", command=self.source_internet)
        # self.lacal_internet.grid(row=0, column=2, pady=(5, 5), padx=(5, 0))
        
        ##############################################################
        #champ de recherche
        self.recherche = customtkinter.CTkEntry(
            self,
            placeholder_text="Rechercher",
            placeholder_text_color="white",
            border_color="#1492E6"
        ).grid(row=4, column=1, columnspan=2, padx=(20, 0), pady=(5, 5), sticky="nsew")
        
        ##################################################
        #bouton pour valider la recher
        self.btn_recherche = customtkinter.CTkOptionMenu(
            master=self.bar_config,
            # fg_color="transparent", 
            # border_width=2, 
            # text_color=("#2C24FE"),
            # text="Exporter",
            values=[
                "Exporter", "Exel", "PDF", "Word"
            ],
            # command=lambda:self.etat(),
            command=lambda eta : self.etat(eta)
            # font=customtkinter.CTkFont(size=18, weight="bold")
        ).grid(row=0, column=1, padx=(20, 20), pady=(10, 10), sticky="nsew")
        
        
        
        #################################################
        #entre des donne
        
        
        
        self.donne = customtkinter.CTkFrame(
            master=self,
            # width=250
        ).grid(row=0, column=1, padx=(0, 0), pady=(0, 0), sticky="nsew")
        
        ######################################
        #onglete des donnes
        self.navigation_donne = customtkinter.CTkFrame(
            self.donne,
            height=50,
            width=250,
        )

        self.navigation_donne.grid(row=0, column=1, sticky="n")
        self.navigation_donne.grid_columnconfigure(1, weight=1)
        self.navigation_donne.grid_columnconfigure((2, 3), weight=0)
        self.navigation_donne.grid_rowconfigure((0, 1, 2), weight=1)

        
        self.severite = customtkinter.CTkLabel(master=self.navigation_donne, text="N_*", width=20).grid(row=0, column=2,padx=(50, 5), pady=(5, 5),sticky="nsew")
        self.severite = customtkinter.CTkLabel(master=self.navigation_donne, text="Name", width=150,).grid(row=0, column=3,padx=(10, 5), pady=(5, 5),sticky="nsew")
        self.alrarm = customtkinter.CTkLabel(master=self.navigation_donne, text="Alarm ID",width=150).grid(row=0, column=4,padx=(10, 5), pady=(5, 5),sticky="nsew")
        self.o_time = customtkinter.CTkLabel(master=self.navigation_donne, text="Occurrence Time",width=150).grid(row=0, column=5,padx=(15, 5), pady=(5, 5),sticky="nsew")
        self.f_ocure = customtkinter.CTkLabel(master=self.navigation_donne, text="Time",width=30).grid(row=0, column=7,padx=(10, 60), pady=(5, 5),sticky="nsew")
        self.l_ocure = customtkinter.CTkLabel(master=self.navigation_donne, text="Last Occurred",width=150).grid(row=0, column=6,padx=(10, 5), pady=(5, 5),sticky="nsew")
        
        

        ###################################
            
            
        #liste de donnes
        self.datas = customtkinter.CTkScrollableFrame(master=self.donne, width=300)
        self.datas.grid(row=0, column=1, padx=(0, 0), pady=(30, 0), sticky="nsew")
        self.datas.grid_columnconfigure(0, weight=1)
        
        self.file = []
        self.code = []
        
        
        ###################################
        #charger les dernier donnees
        try:
            if os.path.exists("filted.csv"):
                # os.remove("filted_data.xlsx")
                for widget in self.datas.winfo_children():
                    widget.destroy()
                    
                data_csv = pd.read_csv("filted.csv")   
                file = data_csv.iloc[0:data_csv.index.stop] # type: ignore
                if self.code != []:
                    self.code = []
                for i in range(file.index.stop):
                    self.code.append(file.loc[i, "MO Name"][:3])
                for i in file.index:
                    self.vue_frame = customtkinter.CTkFrame(master=self.datas, border_width=1, corner_radius=1, width=300, border_color="#1492E6")
                    self.vue_frame.grid(row=i, column=0, padx=5, pady=(0, 10), sticky="nsew")
                    self.vue_frame.grid_columnconfigure(1, weight=1)
                    self.vue_frame.grid_columnconfigure((2, 3), weight=0)
                    self.vue_frame.grid_rowconfigure((0, 1, 2), weight=1)
           
            
                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{i}",width=45)
                    vue.grid(row=i, column=0, padx=(15, 5), pady=(5, 5), sticky="nsew")

                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Severity"]}",width=45)
                    vue.grid(row=i, column=2, padx=(15, 5), pady=(5, 5), sticky="nsew")
                    self.file.append(vue)
            
                    text = file.loc[i,"MO Name"]
                    text = text[7:]
                    vues = customtkinter.CTkLabel(master=self.vue_frame, text=f"{text}",width=45)
                    vues.grid(row=i, column=1, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    vues.configure(text=text)
                    self.file.append(vues)
            
                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Occurrence Times"]}", width=45)
                    vue.grid(row=i, column=4, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    self.file.append(vue)
            
                    # date1_str = file.loc[i,"First Occurred (NT)"]
                    date1 = datetime.now()
                    date2_str = file.loc[i,"Last Occurred (NT)"]
            
                    # date1 = datetime.strptime(date1_str, "%Y-%m-%d %H:%M:%S")
                    date2 = datetime.strptime(date2_str, "%Y-%m-%d %H:%M:%S")
                    difference = date1 - date2
    
            
                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{difference}", width=45)
                    vue.grid(row=i, column=6, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    self.file.append(vue)
            

                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Last Occurred (NT)"]}", width=45)
                    vue.grid(row=i, column=5, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    self.file.append(vue)
            
                    temps = time.time() - temps

       
            
                self.diagrame = customtkinter.CTkLabel(master=self.config, text="Souce").grid(row=1, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{self.type}", text_color="red").grid(row=2, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text="Duree").grid(row=3, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{int(temps)} Seconde", text_color="red").grid(row=4, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text="Ce").grid(row=5, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{time.strftime("%A  %d/%B/%Y")}", text_color="red").grid(row=6, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text="A").grid(row=7, column=0)
                self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{time.strftime("%H/%I/%S")}", text_color="red").grid(row=8, column=0)

            
        except Exception as e:
            print(e)
            
            

            
        
        
        
        #################################################
        #type de statistique
        #positionner sous form de bar de navigation
        self.stat = customtkinter.CTkTabview(
            master=self,
            width=250,
            height=200,
            # fg_color="red"
        )
        self.stat.grid(row=1, column=1, padx=(10, 0), pady=(0, 0), sticky="nsew")
        self.stat.add("Histographe")#ajouter la bar diagrame
        self.stat.add("Graphe")#ajouter la bar ggraphe
        self.stat.add("Cartographie")#ajouter la bar carte
        self.stat.tab("Histographe").grid_columnconfigure(0, weight=1) #creer la bar diagrame dans la bar stat
        self.stat.tab("Graphe").grid_columnconfigure(0, weight=1)#creer la bar graphe dans la bar stat
        self.stat.tab("Cartographie").grid_columnconfigure(0, weight=1)#creer la bar carte dans la bar stat
        #bar diagrame
        self.diagrame = customtkinter.CTkFrame(
            self.stat.tab("Histographe"),

        ).grid(row=0, column=0,)
        
        
        #bar graphe
        self.diagrame = customtkinter.CTkFrame(
            self.stat.tab("Graphe"),height=200
            # dynamic_resizing=False
        ).grid(row=0, column=0)
        #bar cartographie
        self.diagrame = customtkinter.CTkFrame(
            self.stat.tab("Cartographie"),
            # dynamic_resizing=False
        ).grid(row=0, column=0)
        
        ################################################
        #panneau de confiuration
        self.config = customtkinter.CTkTabview(
            master=self,
            width=200,
        )
        self.config.grid(row=0, column=3, padx=(10, 0), pady=(20, 0), sticky="nsew")
        self.config.add("Configuration")#ajouter la bar diagrame
        self.config.tab("Configuration")
        
        ######################################################
        #type d'alarme
        self.config_type_alarmeL = customtkinter.CTkLabel(
            master=self.config.tab("Configuration"),
            text="Type d'alarme"
        ).grid(row=0, column=0,sticky="nsew")
        self.config_type_alarme = customtkinter.CTkOptionMenu(
            master=self.config.tab("Configuration"),
            values=["Majeur", "Mineur", "Critique", "Alert"],
            command=self.choix_type_alarme
        ).grid(row=1, column=0, padx=(25, 0),  pady=(0, 10))
        
        
        #Type de site
        self.config_type_siteL = customtkinter.CTkLabel(
            master=self.config.tab("Configuration"),
            text="Type de site"
        ).grid(row=2, column=0)
        self.config_type_site = customtkinter.CTkOptionMenu(
            master=self.config.tab("Configuration"),
            values=["IHS", "BTS"],
            command=self.choix_type_site
        ).grid(row=3, column=0, padx=(25, 0), pady=(0, 10))
        
        #caracteristique du site
        
        self.config_carac_siteL = customtkinter.CTkLabel(
            master=self.config.tab("Configuration"),
            text="Type de cause"
        ).grid(row=4, column=0)
        self.config_carac_site = customtkinter.CTkOptionMenu(
            master=self.config.tab("Configuration"),
            values=["Transmition", "Energie", "Sinistre"],
            command=self.choix_caract
        ).grid(row=5, column=0, padx=(25, 0))
        

        
        #########################################
        #observation
        self.observation = customtkinter.CTkFrame(
            master=self,
            width=300,
        )
        self.observation.grid(row=0, column=2, padx=(5, 0), pady=(0, 0), sticky="nsew")
        
        self.label = customtkinter.CTkLabel(master=self.observation, text="Observation")
        self.label.grid(row=0, column=0, columnspan=1, padx=20, pady=0, sticky="news")
        
        #######################################
        #diarame reseau
        self.config = customtkinter.CTkScrollableFrame(
            master=self,
            width=150,
            label_text="Diagrame"
        )
        self.config.grid_columnconfigure(0, weight=1)
        self.config.grid(row=1, column=2, padx=(10, 0), pady=(10, 0), sticky="nsew")

        #########################################
        #donne recolter
        self.scrollable_frame = customtkinter.CTkScrollableFrame(master=self, label_text="Statistique", width=150)
        self.scrollable_frame.grid(row=1, column=3, padx=(10, 0), pady=(20, 0), sticky="nsew")
        self.scrollable_frame.grid_columnconfigure(0, weight=1)


        

            
            
            
            
        #################################valeur par defaut
        
        self.fil_path = ''
        self.type = ""
        
        
        
        
    def graph(self):
        f = Figure(figsize=(6.5,3), dpi=90)
        fig = plt.subplots()
        ax = f.add_subplot(111)
        if os.path.exists("filted_data.xlsx"):
            filted_data = pd.read_excel("filted_data.xlsx")
            x = [
                len(filted_data[filted_data["Alarm ID"] == 301]),
                len(filted_data[filted_data["Alarm ID"] == 11522]),
                len(filted_data[filted_data["Alarm ID"] == 1715]),
                len(filted_data[filted_data["Alarm ID"] == 14001]),
                # len(filted_data[filted_data["Alarm ID"] == 26263]),
                # len(filted_data[filted_data["Alarm ID"] == 2658]),
                # len(filted_data[filted_data["Alarm ID"] == 6504]),
                # len(filted_data[filted_data["Alarm ID"] == 21801]),
                # len(filted_data[filted_data["Alarm ID"] == 21825]),
                # len(filted_data[filted_data["Alarm ID"] == 21541]),
                # len(filted_data[filted_data["Alarm ID"] == 22202]),
                # len(filted_data[filted_data["Alarm ID"] == 22214]),
                # len(filted_data[filted_data["Alarm ID"] == 11522])
                ]
            y = [1,2,3,4]
        canvas = FigureCanvasTkAgg(f, master=self.stat.tab("Graphe"))
        
        ax.plot(x, y, c="black")
        ax.set_xlabel("Occurence")
        ax.set_ylabel("Nombre de site")
        ax.set_title("Representaion des alarm par d'occurence")
        canvas.get_tk_widget().grid(sticky="n", row=0, column=0)
        canvas.draw()
        
    def histo(self, file):
        f = Figure(figsize=(6.5,3), dpi=90)
        # fig = plt.subplots()
        ax = f.add_subplot(111)
        
        tab_occ = []
        occure = (file[file["Occurrence Times"] > 1])
        occure.to_excel("occ.xlsx")
        occ = pd.read_excel("occ.xlsx")
        for i in range(occ.index.stop):
            tab_occ.append(int(occ.loc[i, "Occurrence Times"]))
        canvas = FigureCanvasTkAgg(f, master=self.stat.tab("Histographe"))
        
        ax.hist(tab_occ)
        ax.label_outer(True)
        ax.set_xlabel("Nombre d'occurence")
        ax.set_ylabel("Nombre de site")
        ax.set_title("Representaion par nombre d'occurence")
        ax.legend(title="Légende")
        # ax.title("Representaion par nombre d'occurence")
        canvas.get_tk_widget().grid(sticky="n", row=0, column=0)
        canvas.draw()
        
        
    
        
    def choix_caract():
        print("site")    
        
        
    def choix_site(self, choice,):
        try:
            if self.file == []:
                new = customtkinter.CTkToplevel(self)
                new.title("Erreur de fichier")
                new.geometry("340x200")
                new.resizable(FALSE, FALSE)
                new.attributes('-topmost', True)
                new.grab_set()
                # new.grid_columnconfigure(0, weight=1)
                msg = customtkinter.CTkLabel(master=new, text="Aucun fichier n'a ete charger")
                msg.grid(row=0, column=0, padx=(10, 0), pady=(20, 40), sticky="nsew")
                bt1 = customtkinter.CTkButton(master=new, text="Stokage interne", command=lambda:local(), )
                bt1.grid(row=1, column=0, padx=(0, 0))
                bt2 = customtkinter.CTkButton(master=new, text="Internet", command=lambda:internet())
                bt2.grid(row=1, column=1)
                def internet():
                    new.destroy()
                    if os.path.exists("filted_data.xlsx"):
                        os.remove("filted_data.xlsx")
                    self.source_internet()
                    
                def local():
                    new.destroy()
                    if os.path.exists("filted_data.xlsx"):
                        os.remove("filted_data.xlsx")
                    self.source_local()
                    

            else:
                if choice == "Mfoundi":
                    # suprimer tout les enfants contenus dans la frame datas
                    for widget in self.datas.winfo_children():
                        widget.destroy()
                        
                    data_csv = pd.read_csv("filted.csv")
                    data_csv["Site"] = self.code
    #                 data_csv = data_csv.drop([
    #    'Cleared On (NT)', 'Acknowledged On (ST)', 'Cleared By',
    #    'Acknowledged By', 'Clearance Status', 'RRU Name',
    #    'Acknowledgement Status', 'BBU Name', 'eNodeB ID', 'Log Serial Number',
    #    'User Label', 'Equipment Alarm Serial Number', 'Additional Information',
    #    'Maintenance Status'], axis=1)
                    data_csv.to_excel("final.xlsx")
                    file = pd.read_excel("final.xlsx")
                    files = file[file["Site"] == "YDE"]
                    files.to_excel("finale.xlsx")
                    file = pd.DataFrame()
                    file = pd.read_excel("finale.xlsx")
                    # self.datas._create_grid()
                    
                    for i in file.index:
                        vue_frame = customtkinter.CTkFrame(
                            master=self.datas, 
                            border_width=1, 
                            corner_radius=1, 
                            width=250,
                            border_color="#1492E6",
                            # fg_color="#1492E6", 
                            )
                        vue_frame.grid(row=i, column=0, padx=5, pady=(0, 10), sticky="nsew")
                        vue_frame.grid_columnconfigure(1, weight=1)
                        vue_frame.grid_columnconfigure((2, 3), weight=0)
                        vue_frame.grid_rowconfigure((0, 1, 2), weight=1)
                        
                        vue = customtkinter.CTkButton(
                            master=vue_frame, 
                            width=45, 
                            fg_color="#1492E6", 
                            text="Editer!", 
                            command=lambda e=file.loc[i] : self.choix(find = e), 
                            # height=40
                            )
                        vue.grid(row=i, column=6 )
                        
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{i}", width=45)
                        vue.grid(row=i, column=0, padx=(15, 5), pady=(5, 5))
                        
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{file.loc[i,"Severity"]}", width=45)
                        vue.grid(row=i, column=2, padx=(15, 5), pady=(5, 5))
                        self.file.append(vue)
            
                        text = file.loc[i,"MO Name"]
                        text = text[7:]
                        vues = customtkinter.CTkLabel(master=vue_frame, text=f"{text}",width=45 )
                        vues.grid(row=i, column=1, padx=(15, 5), pady=(5, 5),)
                        vues.configure(text=text)
                        # self.file.append(vues)


                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{file.loc[i,"Occurrence Times"]}", width=45)
                        vue.grid(row=i, column=3, padx=(15, 5), pady=(5, 5))
                        # self.file.append(vue)
                        # date1_str = file.loc[i,"First Occurred (NT)"]
                        date1 = datetime.now()
                        date2_str = file.loc[i,"Last Occurred (NT)"]
                        
            
            # date1 = datetime.strptime(date1_str, "%Y-%m-%d %H:%M:%S")
                        date2 = datetime.strptime(date2_str, "%Y-%m-%d %H:%M:%S")
                        difference = date1 - date2
            
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{difference}", width=45)
                        vue.grid(row=i, column=5, padx=(15, 5), pady=(5, 5))
                        # self.file.append(vue)
            
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{file.loc[i,"Last Occurred (NT)"]}", width=45)
                        vue.grid(row=i, column=4, padx=(15, 5), pady=(5, 5))
                        # self.file.append(vue)
                        
                        
                        
                elif choice == "Hors Mfoundi":
                    #dans le cas ou le choix est Hor-Mfoundi
                    alarm = pd.read_excel("alarm.xlsx")
                    hf = pd.DataFrame()
                    code =  []
                    codes = [
    "BMY201","MKN138","YDE198","NEB142","NTU141","BZG210","AYS144",
    "ESK258",'MTE206',"MNB213","NGM199","BJK254","YDE200","MBY252","BFA137",
    ]
                    for i in alarm.index:
                        code.append(alarm.loc[i, "MO Name"][:6])

                    alarm["Sites"] = code

#YDE043 = alarm[alarm["Sites"] == "YDE043"]
                    for i in range(len(codes)):
                        hf = pd.concat([hf, alarm[alarm["Sites"] == codes[i]]], ignore_index=True)
                    hf.to_excel("hf.xlsx")
                        
                    for widget in self.datas.winfo_children():
                        widget.destroy()
                        
                    for i in hf.index:
                        vue_frame = customtkinter.CTkFrame(
                            master=self.datas, 
                            border_width=1, 
                            corner_radius=1, 
                            width=250,
                            border_color="#1492E1" ,
                            # fg_color="#1492E6",
                            )
                        vue_frame.grid(row=i, column=0, padx=5, pady=(0, 10), sticky="nsew")
                        vue_frame.grid_columnconfigure(1, weight=1)
                        vue_frame.grid_columnconfigure((2, 3), weight=0)
                        vue_frame.grid_rowconfigure((0, 1, 2), weight=1)
                        
                                                
                        vue = customtkinter.CTkButton(
                            master=vue_frame, 
                            width=45, 
                            fg_color="#1492E6", 
                            text="Editer!", 
                            command=lambda e=hf.loc[i] : self.choix(find = e), 
                            # height=45
                            )
                        vue.grid(row=i, column=6)
                        
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{i}", width=45)
                        vue.grid(row=i, column=0, padx=(15, 5), pady=(5, 5))
                        
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{hf.loc[i,"Severity"]}", width=45)
                        vue.grid(row=i, column=2, padx=(15, 5), pady=(5, 5))
                        self.file.append(vue)
            
                        text = hf.loc[i,"MO Name"]
                        text = text[7:]
                        vues = customtkinter.CTkLabel(master=vue_frame, text=f"{text}",width=45 )
                        vues.grid(row=i, column=1, padx=(15, 5), pady=(5, 5),)
                        vues.configure(text=text)
                        # self.file.append(vues)
                        # date1_str = file.loc[i,"First Occurred (NT)"]
                        date1 = datetime.now()
                        date2_str = hf.loc[i,"Last Occurred (NT)"]
            
                        # date1 = datetime.strptime(date1_str, "%Y-%m-%d %H:%M:%S")
                        date2 = datetime.strptime(date2_str, "%Y-%m-%d %H:%M:%S")
                        difference = date2 - date1

                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{hf.loc[i,"Occurrence Times"]}", width=45)
                        vue.grid(row=i, column=3, padx=(15, 5), pady=(5, 5))
                        # self.file.append(vue)
            
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{difference}", width=45)
                        vue.grid(row=i, column=5, padx=(15, 5), pady=(5, 5))
                        # self.file.append(vue)
            
                        vue = customtkinter.CTkLabel(master=vue_frame, text=f"{hf.loc[i,"Last Occurred (NT)"]}", width=45)
                        vue.grid(row=i, column=4, padx=(15, 5), pady=(5, 5))
                        # self.file.append(vue)
    
                    
                    
        except Exception as e:
            print(e)
        

    def choix(self, find):
        try:
            detail.detail(self, find)
            
        except Exception as e:
            print(e)
       
    def choix_recap(self, choice):
        if self.file == []:
            try:
                pass
            except EXCEPTION as e:
                print
        else:
            if os.path.exists("filted_data.xlsx"):
                filtre = pd.read_excel("filted_data.xlsx")
                disponible = pd.DataFrame()
                alarm_id = [
                    "Major",
                    "Minor",
                    "Warning"
                    ]
                for i in alarm_id:
                    #dispo.append(filtre[filtre["Severity"] == i])    
                    disponible = pd.concat([disponible, filtre[filtre["Severity"] == i]], ignore_index=True)
                disponible.to_excel("disponible.xlsx")
                if choice == "Mfoundi":
                    dispo.mf(self)
                elif choice == "Hors Mfoundi":
                    dispo.hf(self)
        
    def choix_type_alarme(self, alarm:str):
        
        print("site")
        
        
    def choix_type_site():
        print("recap")
        
        
    def choix_theme(self, new_appearance_mode:str):
        customtkinter.set_appearance_mode(new_appearance_mode)
        
    def choix_echel(self, echelle:str):
        new_scaling_float = int(echelle.replace("%", "")) / 100
        customtkinter.set_widget_scaling(new_scaling_float)
    
    def sup_grid(self):
        self.vue_frame.destroy()
        
    def sources(self, choice):
        if choice == "Stokage Interne":
            self.source_local()
        elif choice == "Internet":
            self.source_internet()
        
    def source_local(self):
        self.type = "Stokage Interne"
        root = Tk()
        root.withdraw()
        self.fil_path = filedialog.askopenfile(filetypes=[("Excel files",  ".xlsx .xls .csv")])
        if os.path.exists("filted_data.xlsx"):
            os.remove("filted_data.xlsx")
        self.actu()
        return self.fil_path, self.type
    
    def source_internet(self):
        self.type = "Internet"
        url = "https://support.huawei.com/enterprise/en/elte-trunking/u2020-m-pid-22461786/software"
        webbrowser.open(url=url)
        self.fil_path = "url"
        return self.fil_path, self.type
    
    def auto(self):
        try:
            if os.path.exists("filted.csv"):
                # os.remove("filted_data.xlsx")
                for widget in self.datas.winfo_children():
                    widget.destroy()
                    
                data_csv = pd.read_csv("filted.csv")   
                file = data_csv.iloc[0:data_csv.index.stop] # type: ignore
                if self.code != []:
                    self.code = []
                for i in range(file.index.stop):
                    self.code.append(file.loc[i, "MO Name"][:3])
                for i in file.index:
                    self.vue_frame = customtkinter.CTkFrame(master=self.datas, border_width=1, corner_radius=1, width=300, border_color="#1492E6")
                    self.vue_frame.grid(row=i, column=0, padx=5, pady=(0, 10), sticky="nsew")
                    self.vue_frame.grid_columnconfigure(1, weight=1)
                    self.vue_frame.grid_columnconfigure((2, 3), weight=0)
                    self.vue_frame.grid_rowconfigure((0, 1, 2), weight=1)
           
            
                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{i}",width=45)
                    vue.grid(row=i, column=0, padx=(15, 5), pady=(5, 5), sticky="nsew")

                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Severity"]}",width=45)
                    vue.grid(row=i, column=2, padx=(15, 5), pady=(5, 5), sticky="nsew")
                    self.file.append(vue)
            
                    text = file.loc[i,"MO Name"]
                    text = text[7:]
                    vues = customtkinter.CTkLabel(master=self.vue_frame, text=f"{text}",width=45)
                    vues.grid(row=i, column=1, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    vues.configure(text=text)
                    self.file.append(vues)
            
                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Occurrence Times"]}", width=45)
                    vue.grid(row=i, column=4, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    self.file.append(vue)
            
                    # date1_str = file.loc[i,"First Occurred (NT)"]
                    date1 = datetime.now()
                    date2_str = file.loc[i,"Last Occurred (NT)"]
            
                    # date1 = datetime.strptime(date1_str, "%Y-%m-%d %H:%M:%S")
                    date2 = datetime.strptime(date2_str, "%Y-%m-%d %H:%M:%S")
                    difference = date1 - date2
    
            
                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{difference}", width=45)
                    vue.grid(row=i, column=6, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    self.file.append(vue)
            

                    vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Last Occurred (NT)"]}", width=45)
                    vue.grid(row=i, column=5, padx=(15, 5), pady=(5, 5),sticky="nsew")
                    self.file.append(vue)
            
                    
        except Exception as e:
            print(e)
        
    def etat(self, choice):
        eta = choice
        # pdf = canvas.Canvas("exemple.pdf",pagesize=letter)
        # largeur, hauteur = letter
        # pdf.drawString(100, hauteur - 100, f"ETAT DU RESEAU MOBILE DANS LA REGION DU CENTRE CE {datetime.now().strftime("%d-%m-%Y A %H:%M:%S")}")
        # pdf.rect(50, hauteur - 150, 200, 100)
        # pdf.save()
        if os.path.exists("etat.xlsx"):
            try:

                if choice == "PDF":
                    pdf = canvas.Canvas("etat.pdf",pagesize=letter)
                    largeur, hauteur = letter
                    pdf.drawString(100, hauteur - 100, f"ETAT DU RESEAU MOBILE DANS LA REGION DU CENTRE CE {datetime.now().strftime("%d-%m-%Y A %H:%M:%S")}")
                    pdf.rect(50, hauteur - 150, 200, 100)
                    root = Tk()
                    root.withdraw()
                    path = filedialog.askdirectory()
                    pdf.save()
                    pass
                
                elif choice == "Exel":
                    data = pd.read_excel("etat.xlsx")
                    doc = Document()
                    doc.add_heading(f"ETAT DU RESEAU MOBILE DANS LA REGION DU CENTRE CE {datetime.now().strftime("%d-%m-%Y A %H:%M:%S")}", 0)
                    tab = doc.add_table(rows=1, cols=len(data.columns), style="Table Grid")
                    
                    hdr_cells = tab.rows[0].cells 
                    for i, column_name in enumerate(data.columns): 
                        hdr_cells[i].text = column_name
                        
                    for index, row in data.iterrows(): 
                        row_cells = tab.add_row().cells 
                        for i, column_name in enumerate(data.columns): 
                            row_cells[i].text = str(row[column_name])
                            
                    doc.save("Etat.docx")
                    document = Document("Etat.docx")
                    
                    root = Tk()
                    root.withdraw()
                    path = filedialog.askdirectory()

                elif choice == "Word":
                    data = pd.read_excel("etat.xlsx")
                    doc = Document()
                    doc.add_heading(f"ETAT DU RESEAU MOBILE DANS LA REGION DU CENTRE CE {datetime.now().strftime("%d-%m-%Y A %H:%M:%S")}", 0)
                    tab = doc.add_table(rows=1, cols=len(data.columns), style="Table Grid")
                    

                    
                    hdr_cells = tab.rows[0].cells 
                    for i, column_name in enumerate(data.columns): 
                        hdr_cells[i].text = column_name
                        
                    for index, row in data.iterrows(): 
                        row_cells = tab.add_row().cells 
                        for i, column_name in enumerate(data.columns): 
                            row_cells[i].text = str(row[column_name])
                    
                    root = Tk()
                    root.withdraw()
                    path = filedialog.askdirectory()
                    doc.save(f"{path}/Etat.docx")
                
            except Exception as e:
                print(e)
                
        else:
            print("chemin non touver")
                

        
            
        

        
        
    def actu(self):
        temps = 0
        temps = time.time()
        try:
            url = self.fil_path
            data = pd.read_excel(url.name)
            data = data.iloc[4:]
            data.columns  = data.iloc[0]
            data = data.iloc[1:]
            if os.path.exists("filted_data.xlsx"):
                filted_data = pd.read_excel("filted_data.xlsx")
                filted_csv = filted_data[filted_data["Alarm ID"] == 301]
                filted_csv.to_csv("filted.csv", index=False)
                filted_csv.to_excel("alarm.xlsx")
                data_csv = pd.read_csv("filted.csv")
            else:
                data.to_excel("filted_data.xlsx")
                filted_data = pd.read_excel("filted_data.xlsx")
                filted_csv = filted_data[filted_data["Alarm ID"] == 301]
                filted_csv.to_csv("filted.csv", index=False)
                filted_csv.to_excel("alarm.xlsx")
                data_csv = pd.read_csv("filted.csv")
            
            file = data_csv.iloc[0:data_csv.index.stop]
            if self.code != []:
                self.code = []
            for i in range(file.index.stop):
                self.code.append(file.loc[i, "MO Name"][:3])
            # file = file.drop([
    #    'Cleared On (NT)', 'Acknowledged On (ST)', 'Cleared By',
    #    'Acknowledged By', 'Clearance Status', 'RRU Name',
    #    'Acknowledgement Status', 'BBU Name', 'eNodeB ID', 'Log Serial Number',
    #    'User Label', 'Equipment Alarm Serial Number', 'Additional Information',
    #    'Maintenance Status'], axis=1)
            self.graph()
            self.histo(file)
        except:
            print("err")
        # suprimer tout les enfants contenus dans la frame datas
        for widget in self.datas.winfo_children():
            widget.destroy()
            
        for i in file.index:
            self.vue_frame = customtkinter.CTkFrame(master=self.datas, border_width=1, corner_radius=1, width=300, border_color="#1492E6")
            self.vue_frame.grid(row=i, column=0, padx=5, pady=(0, 10), sticky="nsew")
            self.vue_frame.grid_columnconfigure(1, weight=1)
            self.vue_frame.grid_columnconfigure((2, 3), weight=0)
            self.vue_frame.grid_rowconfigure((0, 1, 2), weight=1)
            
            
           
            
            vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{i}",width=45)
            vue.grid(row=i, column=0, padx=(15, 5), pady=(5, 5), sticky="nsew")

            vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Severity"]}",width=45)
            vue.grid(row=i, column=2, padx=(15, 5), pady=(5, 5), sticky="nsew")
            self.file.append(vue)
            
            text = file.loc[i,"MO Name"]
            text = text[7:]
            vues = customtkinter.CTkLabel(master=self.vue_frame, text=f"{text}",width=45)
            vues.grid(row=i, column=1, padx=(15, 5), pady=(5, 5),sticky="nsew")
            vues.configure(text=text)
            self.file.append(vues)
            
            vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Occurrence Times"]}", width=45)
            vue.grid(row=i, column=4, padx=(15, 5), pady=(5, 5),sticky="nsew")
            self.file.append(vue)
            
            # date1_str = file.loc[i,"First Occurred (NT)"]
            date1 = datetime.now()
            date2_str = file.loc[i,"Last Occurred (NT)"]
            
            # date1 = datetime.strptime(date1_str, "%Y-%m-%d %H:%M:%S")
            date2 = datetime.strptime(date2_str, "%Y-%m-%d %H:%M:%S")
            difference = date1 - date2
            
    
            
            vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{difference}", width=45)
            vue.grid(row=i, column=6, padx=(15, 5), pady=(5, 5),sticky="nsew")
            self.file.append(vue)
            
            
            vue = customtkinter.CTkLabel(master=self.vue_frame, text=f"{file.loc[i,"Last Occurred (NT)"]}", width=45)
            vue.grid(row=i, column=5, padx=(15, 5), pady=(5, 5),sticky="nsew")
            self.file.append(vue)
            
        temps = time.time() - temps
        # # periode
        # periode = (time.strftime("%d/%m/%Y/ %H/%I"))
        # try:
        #     with open("periode.txt", "a+") as file:
        #         file.write(f"{periode}\n")
        #         file.close()    
        # except:
        #     print("err+++++++++++++++++++++++++++++++++++++++++")
       
            
        self.diagrame = customtkinter.CTkLabel(master=self.config, text="Souce").grid(row=1, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{self.type}", text_color="red").grid(row=2, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text="Duree").grid(row=3, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{int(temps)} Seconde", text_color="red").grid(row=4, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text="Ce").grid(row=5, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{time.strftime("%A  %d/%B/%Y")}", text_color="red").grid(row=6, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text="A").grid(row=7, column=0)
        self.diagrame = customtkinter.CTkLabel(master=self.config, text=f"{time.strftime("%H/%I/%S")}", text_color="red").grid(row=8, column=0)

        
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f"Site critique: {file.index.stop}")
        switch.grid(row=0, column=0, padx=10, pady=(0, 20))
        
        
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f" occurenc eneleve: {max(filted_data["Occurrence Times"])}")
        switch.grid(row=1, column=0, padx=10, pady=(0, 20))
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f"Plus petite occurence: {min(filted_data["Occurrence Times"])}")
        switch.grid(row=2, column=0, padx=10, pady=(0, 20))
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f"Site deconnecter: {len(filted_data[filted_data["Name"] == "NE Is Disconnected"])}")
        switch.grid(row=3, column=0, padx=10, pady=(0, 20))
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f" Site hor-service: {len(filted_data[filted_data["Name"] == "Equipment Out of Service"])}")
        switch.grid(row=4, column=0, padx=10, pady=(0, 20)) 
        
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f" Site Majeur: {len(filted_data[filted_data["Severity"] == "Major"])}")
        switch.grid(row=5, column=0, padx=10, pady=(0, 20)) 
        
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f" Site Mineur: {len(filted_data[filted_data["Severity"] == "Minor"])}")
        switch.grid(row=6, column=0, padx=10, pady=(0, 20)) 
        
        switch = customtkinter.CTkLabel(master=self.scrollable_frame, text=f" Site Alert: {len(filted_data[filted_data["Severity"] == "Warning"])}")
        switch.grid(row=5, column=0, padx=10, pady=(0, 20)) 
                       
if __name__== "__main__":
    app = App()
    app.mainloop()